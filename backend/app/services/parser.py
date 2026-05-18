import io
import time
from typing import List

from fastapi import UploadFile

from app.config import settings
from app.models.schemas import ParsedDocument


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}

# Common encodings to try for plain text, in order of preference
TEXT_ENCODINGS = ["utf-8", "gbk", "gb2312", "gb18030"]


class DocumentParseError(Exception):
    """Raised when document parsing fails."""


class DocumentParser:
    """Parses uploaded documents (PDF, DOCX, TXT) into plain text.

    All parsing is done in-memory without creating temporary files.
    """

    @staticmethod
    def _validate_extension(filename: str) -> None:
        """Raise DocumentParseError if the file extension is not allowed."""
        lower = filename.lower()
        for ext in ALLOWED_EXTENSIONS:
            if lower.endswith(ext):
                return
        raise DocumentParseError(
            f"Unsupported file type: {filename!r}. "
            f"Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    @staticmethod
    def _validate_size(size: int) -> None:
        """Raise DocumentParseError if the file exceeds the size limit."""
        limit = settings.max_file_size_bytes
        if size > limit:
            raise DocumentParseError(
                f"File too large: {size / 1024 / 1024:.1f} MB exceeds "
                f"the {settings.MAX_FILE_SIZE_MB} MB limit"
            )

    async def parse(self, file: UploadFile) -> ParsedDocument:
        """Parse a single uploaded file and return a ParsedDocument.

        Raises DocumentParseError on failure.
        """
        t0 = time.perf_counter()
        filename = file.filename or "unknown"

        # Read the full file content into memory
        content = await file.read()
        self._validate_size(len(content))
        self._validate_extension(filename)

        lower = filename.lower()
        if lower.endswith(".pdf"):
            text = self._parse_pdf(content)
        elif lower.endswith(".docx"):
            text = self._parse_docx(content)
        else:
            text = self._parse_txt(content)

        elapsed_ms = (time.perf_counter() - t0) * 1000
        return ParsedDocument(
            filename=filename,
            text=text,
            char_count=len(text),
            parse_time_ms=round(elapsed_ms, 2),
        )

    def _parse_pdf(self, content: bytes) -> str:
        """Extract text from a PDF using PyMuPDF."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise DocumentParseError(
                "PyMuPDF is not installed. Install with: pip install PyMuPDF"
            )

        try:
            doc = fitz.open(stream=content, filetype="pdf")
        except Exception as exc:
            raise DocumentParseError(f"Failed to open PDF: {exc}")

        pages: List[str] = []
        try:
            for page in doc:
                pages.append(page.get_text())
        finally:
            doc.close()

        return "\n\n".join(pages)

    def _parse_docx(self, content: bytes) -> str:
        """Extract text from a DOCX file using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise DocumentParseError(
                "python-docx is not installed. Install with: pip install python-docx"
            )

        try:
            doc = Document(io.BytesIO(content))
        except Exception as exc:
            raise DocumentParseError(f"Failed to open DOCX: {exc}")

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _parse_txt(self, content: bytes) -> str:
        """Decode plain text with encoding fallbacks."""
        errors_list: List[str] = []
        for encoding in TEXT_ENCODINGS:
            try:
                return content.decode(encoding)
            except (UnicodeDecodeError, LookupError) as exc:
                errors_list.append(f"{encoding}: {exc}")

        raise DocumentParseError(
            f"Failed to decode text file with any supported encoding. "
            f"Errors: {'; '.join(errors_list)}"
        )

    async def parse_batch(self, files: List[UploadFile]) -> tuple[List[ParsedDocument], List[dict]]:
        """Parse multiple files, collecting results and non-fatal errors."""
        results: List[ParsedDocument] = []
        errors: List[dict] = []
        for file in files:
            try:
                result = await self.parse(file)
                results.append(result)
            except DocumentParseError as exc:
                errors.append({
                    "filename": file.filename or "unknown",
                    "error": str(exc),
                })
        return results, errors


# Module-level singleton instance
document_parser = DocumentParser()
